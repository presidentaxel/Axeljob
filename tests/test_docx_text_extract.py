"""AXE-327 — extraction DOCX (tables/runs) + refus .doc legacy."""

from __future__ import annotations

import unittest
from io import BytesIO
from types import SimpleNamespace
from unittest.mock import MagicMock, patch

from fastapi import HTTPException

from backend import main
from backend.services.docx_text_extract import (
    DOC_LEGACY_REFUSAL_DETAIL,
    extract_text_from_docx_bytes,
    is_docx_upload,
    is_legacy_doc,
)


def _docx_with_table_and_soft_break() -> bytes:
    from docx import Document
    from docx.oxml import OxmlElement

    doc = Document()
    doc.add_paragraph("Avant tableau")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Poste"
    table.cell(0, 1).text = "Periode"
    table.cell(1, 0).text = "Dev — AxeL"
    table.cell(1, 1).text = "2024"
    doc.add_paragraph("Apres tableau")
    p = doc.add_paragraph()
    run = p.add_run("SQL")
    run._r.append(OxmlElement("w:br"))
    p.add_run("Python")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestLegacyDocDetection(unittest.TestCase):
    def test_extension_doc(self):
        self.assertTrue(is_legacy_doc(filename="cv.doc"))
        self.assertFalse(is_legacy_doc(filename="cv.docx"))

    def test_content_type_msword(self):
        self.assertTrue(is_legacy_doc(content_type="application/msword"))

    def test_ole_magic(self):
        ole = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 32
        self.assertTrue(is_legacy_doc(file_bytes=ole))

    def test_docx_not_legacy(self):
        self.assertFalse(
            is_legacy_doc(
                filename="cv.docx",
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                file_bytes=b"PK\x03\x04fake",
            )
        )


class TestDocxUploadDetection(unittest.TestCase):
    def test_extension_and_ctype(self):
        self.assertTrue(is_docx_upload(filename="x.docx"))
        self.assertTrue(
            is_docx_upload(
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        )
        self.assertFalse(is_docx_upload(filename="x.doc"))


class TestExtractDocxTablesAndRuns(unittest.TestCase):
    def test_includes_table_cells_in_body_order(self):
        text = extract_text_from_docx_bytes(_docx_with_table_and_soft_break())
        self.assertIn("Avant tableau", text)
        self.assertIn("Poste | Periode", text)
        self.assertIn("Dev — AxeL | 2024", text)
        self.assertIn("Apres tableau", text)
        # Soft break → newline between runs
        self.assertIn("SQL\nPython", text)
        # Table content must not be lost (paragraphs-only extraction would miss it)
        before_table_only = "Avant tableau"
        self.assertNotEqual(text.strip(), before_table_only)


class TestApiCvImportRefusesDoc(unittest.TestCase):
    def test_upload_doc_returns_400_with_convert_message(self):
        req = SimpleNamespace(headers={}, state=SimpleNamespace())
        upload = MagicMock()
        upload.filename = "ancien.doc"
        upload.content_type = "application/msword"
        upload.file = BytesIO(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 64)

        with (
            patch.object(main, "_require_user_id"),
            patch.object(main, "_get_user_id", return_value="user_test"),
            patch.object(main, "check_rate_limit"),
        ):
            with self.assertRaises(HTTPException) as ctx:
                main.api_cv_import_file(req, upload)
        self.assertEqual(ctx.exception.status_code, 400)
        self.assertEqual(ctx.exception.detail, DOC_LEGACY_REFUSAL_DETAIL)
        self.assertIn(".docx", ctx.exception.detail)


if __name__ == "__main__":
    unittest.main()
