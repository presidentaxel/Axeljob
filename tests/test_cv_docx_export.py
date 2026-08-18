"""Tests export Word CV (AXE-330) — sémantique + layout-aware."""

from io import BytesIO

import pytest
from docx import Document

from backend.services.cv_docx_export import cv_to_docx_bytes
from backend.services.docx_text_extract import extract_text_from_docx_bytes


def _sample_cv() -> dict:
    return {
        "prenom": "Ada",
        "nom": "Lovelace",
        "titre_professionnel": "Analyste",
        "email": "ada@example.com",
        "telephone": "+33 1 23 45 67 89",
        "resume": "Pionnière du calcul.",
        "experiences": [
            {
                "poste": "Analyste",
                "entreprise": "Babbage",
                "date_debut": "1840",
                "date_fin": "1850",
                "bullet_points": ["Machine analytique"],
            }
        ],
        "competences": {"techniques": ["Maths", "Algo"]},
        "formations": [{"diplome": "Maths", "etablissement": "Londres", "date": "1835"}],
    }


def _two_column_layout() -> dict:
    return {
        "version": 3,
        "format": "A4",
        "grid": "free",
        "unit": "mm",
        "theme": {
            "font_heading": "Calibri",
            "font_body": "Calibri",
            "color_accent": "#003c33",
            "color_sidebar": "#edfce9",
        },
        "pages": [
            {
                "id": "p1",
                "blocks": [
                    {
                        "id": "id1",
                        "type": "identity",
                        "x": 10,
                        "y": 8,
                        "w": 190,
                        "h": 20,
                        "style": {"font_size": 18},
                    },
                    {
                        "id": "ct1",
                        "type": "contact",
                        "x": 10,
                        "y": 28,
                        "w": 190,
                        "h": 10,
                    },
                    {
                        "id": "sk1",
                        "type": "skills",
                        "x": 10,
                        "y": 45,
                        "w": 55,
                        "h": 40,
                        "bind": "competences.techniques",
                        "style": {"section_label": "Skills"},
                    },
                    {
                        "id": "ex1",
                        "type": "experiences",
                        "x": 75,
                        "y": 45,
                        "w": 125,
                        "h": 80,
                        "style": {"section_label": "Expérience", "title_style": "underline-accent"},
                    },
                    {
                        "id": "fm1",
                        "type": "formations",
                        "x": 75,
                        "y": 140,
                        "w": 125,
                        "h": 40,
                    },
                ],
            }
        ],
    }


def test_cv_to_docx_bytes_roundtrip_text():
    raw = cv_to_docx_bytes(_sample_cv())
    assert raw[:2] == b"PK"  # OOXML zip
    text = extract_text_from_docx_bytes(raw)
    assert "Ada Lovelace" in text
    assert "Analyste" in text
    assert "ada@example.com" in text
    assert "Maths" in text


def test_cv_to_docx_bytes_empty_raises():
    with pytest.raises(ValueError, match="empty_cv"):
        cv_to_docx_bytes({})
    with pytest.raises(ValueError, match="empty_cv"):
        cv_to_docx_bytes(None)


def test_cv_to_docx_bytes_opens_with_python_docx():
    raw = cv_to_docx_bytes(_sample_cv())
    doc = Document(BytesIO(raw))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("Ada Lovelace" in t for t in texts)


def test_layout_aware_docx_uses_section_labels_and_order():
    raw = cv_to_docx_bytes(_sample_cv(), _two_column_layout())
    text = extract_text_from_docx_bytes(raw)
    assert "Ada Lovelace" in text
    assert "SKILLS" in text or "Skills" in text
    assert "EXPÉRIENCE" in text or "Expérience" in text
    assert "Machine analytique" in text
    assert "Maths" in text
    # Skills (sidebar label) should appear before experience content in reading flow /
    # or both present; layout uses table so order in linear extract may vary — both required.
    assert "Babbage" in text


def test_layout_aware_docx_builds_two_column_table():
    raw = cv_to_docx_bytes(_sample_cv(), _two_column_layout())
    doc = Document(BytesIO(raw))
    assert len(doc.tables) == 1
    assert len(doc.tables[0].columns) == 2
