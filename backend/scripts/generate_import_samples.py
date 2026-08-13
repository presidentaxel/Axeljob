#!/usr/bin/env python3
"""Genere les fixtures anonymisees d'import CV (AXE-41).

Usage (depuis la racine du repo) :

    PYTHONPATH=. python backend/scripts/generate_import_samples.py

Ecrase ``tests/fixtures/import_samples/*.pdf|docx``.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

REPO = Path(__file__).resolve().parents[2]
OUT = REPO / "tests" / "fixtures" / "import_samples"


def _fitz():
    """Import PyMuPDF compatible (pymupdf moderne ou alias fitz)."""
    try:
        import pymupdf as fitz
    except ImportError:
        import fitz
    return fitz


def _pdf_single_column() -> bytes:
    fitz = _fitz()

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 50
    lines = [
        ("Camille Durand", 18),
        ("Product Manager", 12),
        ("camille.durand@example.fr  |  +33 6 11 22 33 44  |  Lyon", 10),
        ("", 10),
        ("Profil", 13),
        (
            "PM confirmee, 6 ans d'experience produit B2B SaaS. "
            "Specialisee discovery et priorisation.",
            10,
        ),
        ("", 10),
        ("Experience professionnelle", 13),
        ("Product Manager — NovaSoft (01/2021 - Aujourd'hui)", 10),
        ("- Lancement de 3 features majeures (+18% retention).", 10),
        ("- Pilotage roadmap trimestrielle avec engineering.", 10),
        ("Junior PM — Atelier Data (06/2018 - 12/2020)", 10),
        ("- Discovery utilisateurs et backlog grooming.", 10),
        ("", 10),
        ("Formation", 13),
        ("Master Management — Universite Lyon (2018)", 10),
        ("", 10),
        ("Competences", 13),
        ("Product discovery, SQL, Figma, Agile, Notion", 10),
        ("", 10),
        ("Langues", 13),
        ("Francais (Natif), Anglais (C1)", 10),
    ]
    for text, size in lines:
        if text:
            page.insert_text((50, y), text, fontsize=size, color=(0.1, 0.1, 0.15))
        y += size + 6
    data = doc.tobytes()
    doc.close()
    return data


def _pdf_sidebar() -> bytes:
    fitz = _fitz()

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.draw_rect(fitz.Rect(0, 0, 170, 842), color=None, fill=(0.07, 0.18, 0.28))
    page.insert_text((20, 50), "Alex Martin", fontsize=14, color=(1, 1, 1))
    page.insert_text((20, 72), "Data Analyst", fontsize=10, color=(0.85, 0.9, 1))
    page.insert_text((20, 110), "Contact", fontsize=11, color=(1, 1, 1))
    page.insert_text((20, 130), "alex.martin@example.fr", fontsize=8, color=(0.9, 0.9, 0.9))
    page.insert_text((20, 145), "+33 7 98 76 54 32", fontsize=8, color=(0.9, 0.9, 0.9))
    page.insert_text((20, 160), "Paris", fontsize=8, color=(0.9, 0.9, 0.9))
    page.insert_text((20, 200), "Competences", fontsize=11, color=(1, 1, 1))
    page.insert_text((20, 220), "SQL, Python, Tableau", fontsize=8, color=(0.9, 0.9, 0.9))
    page.insert_text((20, 235), "dbt, Looker, Excel", fontsize=8, color=(0.9, 0.9, 0.9))
    page.insert_text((20, 270), "Langues", fontsize=11, color=(1, 1, 1))
    page.insert_text((20, 290), "Francais, Anglais B2", fontsize=8, color=(0.9, 0.9, 0.9))

    page.insert_text((190, 50), "Profil", fontsize=13, color=(0.1, 0.1, 0.15))
    page.insert_text(
        (190, 70),
        "Data Analyst 5 ans, dashboards et automatisation reporting.",
        fontsize=10,
        color=(0.15, 0.15, 0.2),
    )
    page.insert_text((190, 110), "Experience professionnelle", fontsize=13, color=(0.1, 0.1, 0.15))
    page.insert_text((190, 132), "Data Analyst — Acme (2022 - Aujourd'hui)", fontsize=10)
    page.insert_text((190, 148), "- Migration dashboards vers Tableau.", fontsize=10)
    page.insert_text((190, 164), "- Reduction time-to-insight de 30%.", fontsize=10)
    page.insert_text((190, 190), "Junior Analyst — Beta (2020 - 2021)", fontsize=10)
    page.insert_text((190, 206), "- Suivi KPI marketing.", fontsize=10)
    page.insert_text((190, 250), "Formation", fontsize=13, color=(0.1, 0.1, 0.15))
    page.insert_text((190, 270), "Master Data — Universite Paris (2020)", fontsize=10)
    data = doc.tobytes()
    doc.close()
    return data


def _pdf_dense() -> bytes:
    fitz = _fitz()

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    y = 40
    blocks = [
        ("Samira Benali", 16),
        ("Full-stack Engineer", 11),
        ("samira.benali@example.com | +33 6 55 44 33 22 | Remote", 9),
        ("Profil", 12),
        (
            "Ingenieur full-stack 7 ans : React, FastAPI, Postgres. " "Focus qualite, CI et DX.",
            9,
        ),
        ("Experience professionnelle", 12),
        ("Senior Engineer — Cloudly (2023 - Aujourd'hui)", 9),
        ("- Architecture multi-tenant et observabilite.", 9),
        ("- Mentoring de 3 developpeurs juniors.", 9),
        ("Engineer — StartupX (2019 - 2023)", 9),
        ("- Passage 0→1 de l'API billing.", 9),
        ("- Mise en place Playwright + pytest CI.", 9),
        ("Dev — Agence Web (2017 - 2019)", 9),
        ("- Sites vitrine et boutiques Shopify.", 9),
        ("Projets", 12),
        ("Open-source cv-lint — linter ATS pour CV markdown.", 9),
        ("Formation", 12),
        ("Diplome Ingenieur — INSA (2017)", 9),
        ("Certifications", 12),
        ("AWS Solutions Architect Associate (2024)", 9),
        ("Competences", 12),
        ("TypeScript, Python, React, FastAPI, Postgres, Docker, K8s", 9),
        ("Langues", 12),
        ("Francais natif, Anglais C1, Arabe B1", 9),
    ]
    for text, size in blocks:
        page.insert_text((40, y), text, fontsize=size, color=(0.05, 0.05, 0.08))
        y += size + 5
    data = doc.tobytes()
    doc.close()
    return data


def _docx_single_column() -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("Camille Durand", level=1)
    doc.add_paragraph("Product Manager")
    doc.add_paragraph("camille.durand@example.fr | +33 6 11 22 33 44 | Lyon")
    doc.add_heading("Profil", level=2)
    doc.add_paragraph(
        "PM confirmee, 6 ans d'experience produit B2B SaaS. "
        "Specialisee discovery et priorisation."
    )
    doc.add_heading("Experience professionnelle", level=2)
    doc.add_paragraph("Product Manager — NovaSoft (01/2021 - Aujourd'hui)")
    doc.add_paragraph("Lancement de 3 features majeures (+18% retention).", style="List Bullet")
    doc.add_heading("Formation", level=2)
    doc.add_paragraph("Master Management — Universite Lyon (2018)")
    doc.add_heading("Competences", level=2)
    doc.add_paragraph("Product discovery, SQL, Figma, Agile, Notion")
    doc.add_heading("Langues", level=2)
    doc.add_paragraph("Francais (Natif), Anglais (C1)")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    mapping = {
        "01_single_column.pdf": _pdf_single_column(),
        "02_sidebar.pdf": _pdf_sidebar(),
        "03_dense_multisection.pdf": _pdf_dense(),
        "04_single_column.docx": _docx_single_column(),
    }
    for name, payload in mapping.items():
        path = OUT / name
        path.write_bytes(payload)
        print(f"  wrote {path.name} ({len(payload)} bytes)")
    print(f"Fixtures ecrites dans {OUT}")


if __name__ == "__main__":
    main()
