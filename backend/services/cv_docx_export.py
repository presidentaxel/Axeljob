"""Export Word (.docx) d'un CV (AXE-330).

- Sans layout : document sémantique mono-colonne (ATS).
- Avec layout free-canvas : flux design-aware (ordre de lecture, 1–2 colonnes,
  typo / accents du thème). Word ne peut pas reproduire le placement absolu mm.
"""

from __future__ import annotations

from io import BytesIO
from typing import Any

from backend.services import layout_bindings as bind
from backend.services.cv_text_export import (
    _as_str,
    _join_nonempty,
    _list_skills,
)
from backend.services.layout_renderer import SECTION_LABELS

PAGE_WIDTH_MM = 210.0
COLUMN_GAP_MM = 30.0
FULL_WIDTH_MIN_MM = 140.0
CONTENT_TYPES = frozenset(
    {
        "identity",
        "contact",
        "resume",
        "experiences",
        "formations",
        "certifications",
        "projets",
        "skills",
        "languages",
        "text",
        "title",
    }
)


def _has_content(cv: dict) -> bool:
    if _as_str(cv.get("prenom")) or _as_str(cv.get("nom")):
        return True
    if _as_str(cv.get("titre_professionnel")) or _as_str(cv.get("resume")):
        return True
    if _as_str(cv.get("email")) or _as_str(cv.get("telephone")):
        return True
    for key in ("experiences", "formations", "certifications", "projets"):
        rows = cv.get(key)
        if isinstance(rows, list) and any(isinstance(r, dict) and r for r in rows):
            return True
    competences = cv.get("competences")
    if isinstance(competences, dict) and competences:
        return True
    return False


def _num(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default


def _reading_key(block: dict[str, Any]) -> tuple[float, float]:
    return (_num(block.get("y")), _num(block.get("x")))


def _theme(layout: dict[str, Any] | None) -> dict[str, Any]:
    if not isinstance(layout, dict):
        return {}
    theme = layout.get("theme")
    return theme if isinstance(theme, dict) else {}


def _layout_has_pages(layout: dict | None) -> bool:
    if not isinstance(layout, dict):
        return False
    pages = layout.get("pages")
    return isinstance(pages, list) and any(
        isinstance(p, dict) and isinstance(p.get("blocks"), list) and p.get("blocks") for p in pages
    )


def _iter_page_blocks(layout: dict[str, Any]) -> list[dict[str, Any]]:
    """Blocs de contenu de la première page (export Word = page 1)."""
    pages = layout.get("pages")
    if not isinstance(pages, list) or not pages:
        return []
    page = pages[0] if isinstance(pages[0], dict) else {}
    blocks = page.get("blocks")
    if not isinstance(blocks, list):
        return []
    out: list[dict[str, Any]] = []
    for block in blocks:
        if not isinstance(block, dict):
            continue
        if str(block.get("type") or "") in CONTENT_TYPES:
            out.append(block)
    return out


def _cluster_column_xs(blocks: list[dict[str, Any]]) -> list[float]:
    xs = sorted({round(_num(b.get("x")), 1) for b in blocks})
    if not xs:
        return [0.0]
    clusters = [xs[0]]
    for x in xs[1:]:
        if x - clusters[-1] >= COLUMN_GAP_MM:
            clusters.append(x)
    return clusters[:2]


def _split_header_and_columns(
    blocks: list[dict[str, Any]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]]]:
    """Sépare bandeau pleine largeur vs colonnes gauche/droite."""
    clusters = _cluster_column_xs(blocks)
    if len(clusters) < 2:
        return ([], sorted(blocks, key=_reading_key), [])

    split_x = (clusters[0] + clusters[1]) / 2.0
    header: list[dict[str, Any]] = []
    left: list[dict[str, Any]] = []
    right: list[dict[str, Any]] = []
    for block in blocks:
        x = _num(block.get("x"))
        w = _num(block.get("w"), 20)
        # Bandeau pleine largeur uniquement (pas les blocs sidebar qui dépassent un peu).
        if w >= FULL_WIDTH_MIN_MM:
            header.append(block)
        elif x < split_x:
            left.append(block)
        else:
            right.append(block)
    header.sort(key=_reading_key)
    left.sort(key=_reading_key)
    right.sort(key=_reading_key)
    return (header, left, right)


def _sidebar_fill_hex(layout: dict[str, Any]) -> str | None:
    theme = _theme(layout)
    sidebar = theme.get("color_sidebar")
    if isinstance(sidebar, str) and sidebar.strip().startswith("#"):
        return sidebar.strip()
    pages = layout.get("pages")
    if not isinstance(pages, list) or not pages:
        return None
    page = pages[0] if isinstance(pages[0], dict) else {}
    for block in page.get("blocks") or []:
        if not isinstance(block, dict):
            continue
        if str(block.get("type") or "") != "shape":
            continue
        style = block.get("style") if isinstance(block.get("style"), dict) else {}
        fill = style.get("bg") or style.get("color") or style.get("fill")
        if not (isinstance(fill, str) and fill.strip().startswith("#")):
            continue
        # Bandeau gauche haut → teinte sidebar.
        if _num(block.get("x")) < 40 and _num(block.get("h")) > 80:
            return fill.strip()
    return None


def _parse_hex_color(raw: Any):
    from docx.shared import RGBColor

    if not isinstance(raw, str):
        return None
    s = raw.strip().lstrip("#")
    if len(s) == 3:
        s = "".join(ch * 2 for ch in s)
    if len(s) != 6:
        return None
    try:
        return RGBColor(int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16))
    except ValueError:
        return None


def _set_run_font(
    run: Any,
    *,
    size_pt: float | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    color_hex: str | None = None,
    font_name: str | None = None,
) -> None:
    from docx.shared import Pt

    if font_name:
        run.font.name = font_name
    if size_pt is not None:
        run.font.size = Pt(max(7.0, min(size_pt, 28.0)))
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rgb = _parse_hex_color(color_hex)
    if rgb is not None:
        run.font.color.rgb = rgb


def _add_para(
    container: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    size_pt: float | None = None,
    color_hex: str | None = None,
    font_name: str | None = None,
    space_after_pt: float = 4,
) -> Any:
    from docx.shared import Pt

    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_run_font(
        run,
        size_pt=size_pt,
        bold=bold,
        italic=italic,
        color_hex=color_hex,
        font_name=font_name,
    )
    return p


def _section_title(block: dict[str, Any], key: str, *, default_title: bool) -> str:
    style = block.get("style") if isinstance(block.get("style"), dict) else {}
    custom = str(style.get("section_label") or "").strip()
    if custom:
        return custom
    if default_title:
        return SECTION_LABELS.get(key, key)
    return ""


def _emit_section_heading(
    container: Any,
    title: str,
    theme: dict[str, Any],
    style: dict[str, Any],
) -> None:
    if not title:
        return
    accent = str(style.get("color") or theme.get("color_accent") or "#1e293b").strip() or "#1e293b"
    font = str(style.get("font_family") or theme.get("font_heading") or "").strip() or None
    size = _num(style.get("font_size"), 11) if style.get("font_size") is not None else 11.0
    p = _add_para(
        container,
        title.upper(),
        bold=True,
        size_pt=size,
        color_hex=accent,
        font_name=font,
        space_after_pt=2,
    )
    # Soulignement discret via bordure bas (approx title_style underline).
    title_style = str(style.get("title_style") or "")
    if title_style in {"underline-accent", "underline", "sidebar-bar"} or not title_style:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "4")
        rgb = accent.lstrip("#")
        if len(rgb) == 3:
            rgb = "".join(ch * 2 for ch in rgb)
        if len(rgb) == 6:
            bottom.set(qn("w:color"), rgb.upper())
        pBdr.append(bottom)
        pPr.append(pBdr)


def _body_font(theme: dict[str, Any], style: dict[str, Any]) -> str | None:
    name = str(
        style.get("font_family") or theme.get("font_body") or theme.get("font_heading") or ""
    )
    return name.strip() or None


def _emit_block(container: Any, cv: dict, block: dict[str, Any], theme: dict[str, Any]) -> None:
    btype = str(block.get("type") or "")
    style = block.get("style") if isinstance(block.get("style"), dict) else {}
    binding = block.get("bind")
    limit = block.get("limit")
    body_font = _body_font(theme, style)
    body_size = _num(style.get("font_size"), 10) if style.get("font_size") is not None else 10.0
    body_color = str(style.get("color_body") or style.get("color") or "#1e293b")

    if btype == "identity":
        name = bind.resolve_bound_text(cv, ["prenom", "nom"])
        title = bind.resolve_bound_text(cv, "titre_professionnel")
        heading_font = (
            str(style.get("font_family") or theme.get("font_heading") or "").strip() or None
        )
        if name:
            _add_para(
                container,
                name,
                bold=True,
                size_pt=(
                    _num(style.get("font_size"), 18) if style.get("font_size") is not None else 18
                ),
                color_hex=str(style.get("color") or theme.get("color_accent") or "#0f172a"),
                font_name=heading_font,
                space_after_pt=2,
            )
        if title:
            if style.get("header_layout") == "inline-title" and name:
                # Already separate lines is clearer in Word.
                pass
            _add_para(
                container,
                title,
                size_pt=11,
                color_hex="#475569",
                font_name=body_font,
                space_after_pt=6,
            )
        return

    if btype == "contact":
        parts = [
            bind.resolve_bound_text(cv, "email"),
            bind.resolve_bound_text(cv, "telephone"),
            bind.resolve_bound_text(cv, "linkedin"),
            bind.resolve_bound_text(cv, "ville") or _as_str(cv.get("localisation")),
        ]
        line = _join_nonempty([p for p in parts if p])
        if line:
            _add_para(
                container,
                line,
                size_pt=9,
                color_hex="#334155",
                font_name=body_font,
                space_after_pt=8,
            )
        return

    if btype in {"text", "title"}:
        content = _as_str(block.get("content"))
        if not content:
            return
        is_title = btype == "title"
        _add_para(
            container,
            content,
            bold=is_title or bool(style.get("bold")),
            italic=bool(style.get("italic")),
            size_pt=body_size if style.get("font_size") is not None else (12 if is_title else 10),
            color_hex=str(
                style.get("color") or (theme.get("color_accent") if is_title else body_color)
            ),
            font_name=body_font,
        )
        return

    if btype == "resume":
        text = bind.resolve_bound_text(cv, binding if binding else "resume")
        title = _section_title(block, "resume", default_title=False)
        if title:
            _emit_section_heading(container, title, theme, style)
        if text:
            _add_para(container, text, size_pt=body_size, color_hex=body_color, font_name=body_font)
        return

    if btype == "experiences":
        items = bind.resolve_experiences(cv, limit if isinstance(limit, int | float) else None)
        title = _section_title(block, "experiences", default_title=True)
        _emit_section_heading(container, title, theme, style)
        for exp in items:
            ent = _as_str(exp.get("entreprise"))
            poste = _as_str(exp.get("poste"))
            dates = _join_nonempty(
                [_as_str(exp.get("date_debut")), _as_str(exp.get("date_fin"))],
                sep=" – ",
            )
            lieu = _as_str(exp.get("lieu"))
            head = ent or poste
            if head:
                meta = _join_nonempty([dates, lieu])
                label = f"{head}  {meta}" if meta else head
                _add_para(
                    container,
                    label,
                    bold=True,
                    size_pt=body_size,
                    color_hex=body_color,
                    font_name=body_font,
                    space_after_pt=1,
                )
            if poste and ent:
                _add_para(
                    container,
                    poste,
                    italic=True,
                    size_pt=max(8.0, body_size - 1),
                    color_hex="#64748b",
                    font_name=body_font,
                    space_after_pt=1,
                )
            for bullet in exp.get("bullet_points") or []:
                s = _as_str(bullet)
                if s:
                    _add_para(
                        container,
                        f"• {s}",
                        size_pt=max(8.0, body_size - 1),
                        color_hex=body_color,
                        font_name=body_font,
                        space_after_pt=1,
                    )
        return

    if btype == "formations":
        items = bind.resolve_formations(cv, limit if isinstance(limit, int | float) else None)
        title = _section_title(block, "formations", default_title=True)
        _emit_section_heading(container, title, theme, style)
        for row in items:
            dip = _as_str(row.get("diplome"))
            etab = _as_str(row.get("etablissement"))
            date = _as_str(row.get("date"))
            head = _join_nonempty([dip, etab])
            if head and date:
                head = f"{head} ({date})"
            elif date and not head:
                head = date
            if head:
                _add_para(
                    container,
                    head,
                    bold=True,
                    size_pt=body_size,
                    color_hex=body_color,
                    font_name=body_font,
                )
        return

    if btype == "certifications":
        items = bind.resolve_certifications(cv, limit if isinstance(limit, int | float) else None)
        title = _section_title(block, "certifications", default_title=True)
        _emit_section_heading(container, title, theme, style)
        for row in items:
            line = _join_nonempty(
                [_as_str(row.get("nom")), _as_str(row.get("organisme")), _as_str(row.get("date"))]
            )
            if line:
                _add_para(
                    container, line, size_pt=body_size, color_hex=body_color, font_name=body_font
                )
        return

    if btype == "projets":
        items = bind.resolve_projets(cv, limit if isinstance(limit, int | float) else None)
        title = _section_title(block, "projets", default_title=True)
        _emit_section_heading(container, title, theme, style)
        for row in items:
            nom = _as_str(row.get("nom"))
            desc = _as_str(row.get("description"))
            if nom:
                _add_para(
                    container,
                    nom,
                    bold=True,
                    size_pt=body_size,
                    color_hex=body_color,
                    font_name=body_font,
                    space_after_pt=1,
                )
            if desc:
                _add_para(
                    container, desc, size_pt=body_size, color_hex=body_color, font_name=body_font
                )
        return

    if btype == "skills":
        items = bind.resolve_bound_string_list(cv, binding if binding else "competences.techniques")
        title = _section_title(block, "skills", default_title=False)
        if title:
            _emit_section_heading(container, title, theme, style)
        if items:
            _add_para(
                container,
                ", ".join(items),
                size_pt=body_size,
                color_hex=body_color,
                font_name=body_font,
            )
        return

    if btype == "languages":
        items = bind.resolve_langues(cv)
        title = _section_title(block, "languages", default_title=True)
        _emit_section_heading(container, title, theme, style)
        labels: list[str] = []
        for row in items:
            label = _as_str(row.get("langue"))
            niveau = _as_str(row.get("niveau"))
            if niveau:
                label = f"{label} ({niveau})"
            if label:
                labels.append(label)
        if labels:
            _add_para(
                container,
                ", ".join(labels),
                size_pt=body_size,
                color_hex=body_color,
                font_name=body_font,
            )


def _set_cell_shading(cell: Any, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    rgb = hex_color.strip().lstrip("#")
    if len(rgb) == 3:
        rgb = "".join(ch * 2 for ch in rgb)
    if len(rgb) != 6:
        return
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), rgb.upper())
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _emit_blocks(
    container: Any, cv: dict, blocks: list[dict[str, Any]], theme: dict[str, Any]
) -> None:
    for block in blocks:
        _emit_block(container, cv, block, theme)


def _build_layout_docx(cv: dict, layout: dict[str, Any]) -> bytes:
    from docx import Document
    from docx.shared import Cm, Pt

    doc = Document()
    # Marges un peu plus serrées pour coller au canvas.
    for section in doc.sections:
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    theme = _theme(layout)
    blocks = _iter_page_blocks(layout)
    if not blocks:
        # Layout vide → fallback sémantique via même document path.
        raise ValueError("empty_layout")

    header, left, right = _split_header_and_columns(blocks)
    _emit_blocks(doc, cv, header, theme)

    if left and right:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        left_cell, right_cell = table.rows[0].cells
        # Largeur relative ~ sidebar 35% / main 65% selon positions.
        try:
            left_w = max((_num(b.get("w")) for b in left), default=60)
            right_w = max((_num(b.get("w")) for b in right), default=120)
            total = max(left_w + right_w, 1)
            usable = Cm(17.5)
            left_cell.width = int(usable * (left_w / total))
            right_cell.width = int(usable * (right_w / total))
        except Exception:
            pass

        sidebar_hex = _sidebar_fill_hex(layout)
        # Si la colonne gauche a plus de blocs « sidebar » (skills/languages) et x bas,
        # teinter la cellule gauche.
        left_is_sidebar = sum(1 for b in left if _num(b.get("x")) < 80) >= len(left) / 2
        if sidebar_hex and left_is_sidebar:
            _set_cell_shading(left_cell, sidebar_hex)

        left_cell.text = ""
        right_cell.text = ""
        _emit_blocks(left_cell, cv, left, theme)
        _emit_blocks(right_cell, cv, right, theme)
    else:
        single = left or right
        _emit_blocks(doc, cv, single, theme)

    # Évite un paragraphe vide orphelin excessif.
    style = doc.styles["Normal"]
    style.font.size = Pt(10)
    body_name = str(theme.get("font_body") or theme.get("font_heading") or "").strip()
    if body_name:
        style.font.name = body_name

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_semantic_docx(cv: dict) -> bytes:
    from docx import Document

    doc = Document()

    identity = _join_nonempty(
        [_as_str(cv.get("prenom")), _as_str(cv.get("nom"))],
        sep=" ",
    )
    if identity:
        _add_para(doc, identity, bold=True, size_pt=18, space_after_pt=2)
    titre = _as_str(cv.get("titre_professionnel"))
    if titre:
        _add_para(doc, titre, bold=True, size_pt=11, color_hex="#475569", space_after_pt=4)

    contact = _join_nonempty(
        [
            _as_str(cv.get("email")),
            _as_str(cv.get("telephone")),
            _as_str(cv.get("linkedin")),
            _as_str(cv.get("ville") or cv.get("localisation")),
        ]
    )
    if contact:
        _add_para(doc, contact, size_pt=9, color_hex="#334155", space_after_pt=8)

    resume = _as_str(cv.get("resume"))
    if resume:
        _add_para(doc, "PROFIL", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2)
        _add_para(doc, resume, size_pt=10)

    experiences = [r for r in (cv.get("experiences") or []) if isinstance(r, dict)]
    if experiences:
        _add_para(doc, "EXPÉRIENCES", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2)
        for row in experiences:
            head = _join_nonempty([_as_str(row.get("poste")), _as_str(row.get("entreprise"))])
            dates = _join_nonempty(
                [_as_str(row.get("date_debut")), _as_str(row.get("date_fin"))],
                sep=" – ",
            )
            if head:
                line = head if not dates else f"{head} ({dates})"
                _add_para(doc, line, bold=True, size_pt=10, space_after_pt=1)
            desc = _as_str(row.get("description") or row.get("missions"))
            if desc:
                _add_para(doc, desc, size_pt=9)

    formations = [r for r in (cv.get("formations") or []) if isinstance(r, dict)]
    if formations:
        _add_para(doc, "FORMATIONS", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2)
        for row in formations:
            head = _join_nonempty([_as_str(row.get("diplome")), _as_str(row.get("etablissement"))])
            date = _as_str(row.get("date"))
            if head:
                line = head if not date else f"{head} ({date})"
                _add_para(doc, line, bold=True, size_pt=10)

    competences = cv.get("competences") if isinstance(cv.get("competences"), dict) else {}
    tech = _list_skills(competences.get("techniques"))
    soft = _list_skills(competences.get("logiciels") or competences.get("outils"))
    autres = _list_skills(competences.get("autres"))
    if tech or soft or autres:
        _add_para(doc, "COMPÉTENCES", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2)
        if tech:
            _add_para(doc, "Techniques : " + ", ".join(tech), size_pt=10)
        if soft:
            _add_para(doc, "Outils : " + ", ".join(soft), size_pt=10)
        if autres:
            _add_para(doc, "Autres : " + ", ".join(autres), size_pt=10)

    lang_lines: list[str] = []
    for row in competences.get("langues") or []:
        if isinstance(row, str):
            s = row.strip()
            if s:
                lang_lines.append(s)
            continue
        if isinstance(row, dict):
            s = _join_nonempty([_as_str(row.get("langue")), _as_str(row.get("niveau"))])
            if s:
                lang_lines.append(s)
    if lang_lines:
        _add_para(doc, "LANGUES", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2)
        for line in lang_lines:
            _add_para(doc, line, size_pt=10)

    cert_lines: list[str] = []
    for row in cv.get("certifications") or []:
        if not isinstance(row, dict):
            continue
        s = _join_nonempty(
            [_as_str(row.get("nom")), _as_str(row.get("organisme")), _as_str(row.get("date"))]
        )
        if s:
            cert_lines.append(s)
    if cert_lines:
        _add_para(
            doc, "CERTIFICATIONS", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2
        )
        for line in cert_lines:
            _add_para(doc, line, size_pt=10)

    projets = [r for r in (cv.get("projets") or []) if isinstance(r, dict)]
    if projets:
        _add_para(doc, "PROJETS", bold=True, size_pt=11, color_hex="#1e293b", space_after_pt=2)
        for row in projets:
            head = _as_str(row.get("nom"))
            desc = _as_str(row.get("description"))
            if head:
                _add_para(doc, head, bold=True, size_pt=10, space_after_pt=1)
            if desc:
                _add_para(doc, desc, size_pt=9)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def cv_to_docx_bytes(cv: dict | None, layout: dict | None = None) -> bytes:
    """
    Sérialise le CV en .docx.

    Si ``layout`` free-canvas est fourni, produit un document design-aware
    (ordre spatial, 2 colonnes, thème). Sinon, export sémantique simple.
    Lève ValueError si le CV (et le layout) ne permettent aucun contenu.
    """
    data = cv if isinstance(cv, dict) else {}
    if _layout_has_pages(layout):
        assert isinstance(layout, dict)
        try:
            return _build_layout_docx(data, layout)
        except ValueError:
            pass
    if not _has_content(data):
        raise ValueError("empty_cv")
    return _build_semantic_docx(data)
