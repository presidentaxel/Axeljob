"""Extraction texte DOCX robuste (tables, runs) + detection .doc legacy (AXE-327)."""

from __future__ import annotations

from io import BytesIO

# Message UX : refuse explicite du Word 97-2003 (.doc / OLE).
DOC_LEGACY_REFUSAL_DETAIL = (
    "Le format Word .doc (ancien) n'est pas supporté. "
    "Ouvre le fichier dans Word ou LibreOffice, enregistre-le en .docx, puis réessaie."
)

# Compound File Binary Format (OLE) — signature des .doc legacy.
_OLE_MAGIC = b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
# OOXML (.docx) is a ZIP package.
_ZIP_MAGIC = b"PK"

_DOCX_CONTENT_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
_MSWORD_CONTENT_TYPE = "application/msword"


def _looks_like_ooxml(file_bytes: bytes | None) -> bool:
    return bool(file_bytes) and file_bytes[:2] == _ZIP_MAGIC


def is_legacy_doc(
    *,
    filename: str | None = None,
    content_type: str | None = None,
    file_bytes: bytes | None = None,
) -> bool:
    """True si l'upload ressemble a un Word .doc legacy (a refuser).

    Ne refuse pas un vrai ``.docx`` mal etiquete ``application/msword``
    (extension ``.docx`` ou magic ZIP/OOXML).
    """
    name = (filename or "").strip().lower()
    ctype = (content_type or "").strip().lower()

    # Preuve OOXML → jamais legacy (meme si MIME msword).
    if name.endswith(".docx") or _looks_like_ooxml(file_bytes):
        return False

    if name.endswith(".doc"):
        return True
    if file_bytes and len(file_bytes) >= 8 and file_bytes[:8] == _OLE_MAGIC:
        return True
    # MIME msword sans preuve OOXML → legacy.
    if ctype == _MSWORD_CONTENT_TYPE:
        return True
    return False


def is_docx_upload(
    *,
    filename: str | None = None,
    content_type: str | None = None,
    file_bytes: bytes | None = None,
) -> bool:
    """True si l'upload est un .docx (extension, content-type OOXML, ou ZIP + MIME Word)."""
    name = (filename or "").strip().lower()
    ctype = (content_type or "").strip().lower()
    if name.endswith(".docx"):
        return True
    if ctype == _DOCX_CONTENT_TYPE:
        return True
    # Clients qui envoient encore application/msword pour un vrai .docx.
    if _looks_like_ooxml(file_bytes) and ctype in (
        _MSWORD_CONTENT_TYPE,
        "application/octet-stream",
        "",
    ):
        return True
    return False


def _paragraph_plain_text(paragraph) -> str:
    """Texte d'un paragraphe (inclut hyperliens / soft breaks via python-docx)."""
    # ``paragraph.text`` parcourt aussi les runs sous ``w:hyperlink`` ;
    # un walk de ``paragraph.runs`` seul les omet.
    return (paragraph.text or "").strip()


def _table_to_text(table) -> str:
    """Aplatit un tableau en lignes « cellule | cellule » (ordre des lignes)."""
    lines: list[str] = []
    for row in table.rows:
        seen_tc: set[int] = set()
        cells_out: list[str] = []
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tc:
                continue
            seen_tc.add(tc_id)
            parts: list[str] = []
            for p in cell.paragraphs:
                t = _paragraph_plain_text(p)
                if t:
                    parts.append(t)
            for nested in cell.tables:
                nested_txt = _table_to_text(nested)
                if nested_txt:
                    parts.append(nested_txt)
            cell_txt = " ".join(parts).strip()
            if cell_txt:
                cells_out.append(cell_txt)
        if cells_out:
            lines.append(" | ".join(cells_out))
    return "\n".join(lines)


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extrait le texte d'un .docx en preservant l'ordre corps (paragraphes + tables)."""
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(BytesIO(file_bytes))
    blocks: list[str] = []
    for item in doc.iter_inner_content():
        if isinstance(item, Paragraph):
            text = _paragraph_plain_text(item)
            if text:
                blocks.append(text)
        elif isinstance(item, Table):
            text = _table_to_text(item)
            if text:
                blocks.append(text)
    return "\n".join(blocks).strip()
